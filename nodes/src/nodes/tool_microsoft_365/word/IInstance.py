# =============================================================================
# RocketRide Engine
# =============================================================================
# MIT License
# Copyright (c) 2026 Aparavi Software AG
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in
# all copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.
# =============================================================================

"""
Word tool node instance.

Exposes a docx round-trip over the Microsoft Graph drive content API: read
text, create a new document, append paragraphs, find/replace text, export to
PDF, and check the connection. Write operations require the ``write`` tier.

There is no persisted Word Online editing session — every write downloads
the current ``.docx``, edits it in-process with ``python-docx``, and
re-uploads the whole file with an ``If-Match`` header carrying the eTag read
at download time. If another editor changed the file in between, the upload
gets HTTP 409/412 back from Graph, which ``graph_client.request`` turns into
a ``GraphError`` naming the conflict; the caller should re-read (via
``word_read_text``) and retry rather than silently overwriting.

``python-docx`` is imported lazily inside each tool method (not at module
import time): the dependency is installed at runtime by the suite's
``requirements.txt`` via ``depends()`` in ``IGlobal.beginGlobal``, and a
top-level import would break test collection in environments where it is
not yet installed.

Operational targets (file path/id, new-document path) are always
invoke-time parameters — never node config.
"""

from __future__ import annotations

import urllib.parse

from io import BytesIO

from rocketlib import tool_function

from ai.common.utils import normalize_tool_input, require_str, require_str_list

from .. import graph_client
from ..IInstance import MicrosoftToolInstanceBase
from .client import SERVICE, WORD_CONTENT_TYPE, _seg, clean_item, download_docx, it, request, upload_docx
from .IGlobal import IGlobal


def _replace_in_paragraph(paragraph, find: str, replace: str) -> int:
    """Replace every occurrence of ``find`` with ``replace`` in a paragraph; return the count.

    Single pass, computed from the paragraph's ORIGINAL text
    (``paragraph.text``, the concatenation of every run's text at the
    moment this function is called) — never from text already mutated by
    an earlier pass. Re-scanning mutated text is how a two-pass
    run-then-paragraph approach double-counts and corrupts the result
    whenever ``replace`` itself contains ``find`` as a substring: e.g.
    find='foo', replace='foobar' on 'foo is here' would count the match
    once in a run-level pass, rewrite it to 'foobar is here', then find
    'foo' *again* inside the freshly written 'foobar' on a second pass,
    counting 2 and corrupting the text to 'foobarbar is here'. Scanning
    the original text exactly once avoids that entirely.

    Paragraphs with zero matches are left completely untouched (runs and
    their formatting unchanged). Paragraphs with at least one match have
    their new text written back as:

    - **Single run**: set that run's text directly — formatting is
      unaffected.
    - **Multiple runs**: set the first run's text to the whole new
      paragraph text and blank every other run. The first run's
      formatting wins for the entire merged text — a paragraph made of
      multiple runs (whether or not the match itself spans a run
      boundary) loses its intra-paragraph formatting (bold/italic/etc.
      boundaries) on any replacement. This is a known, documented
      trade-off of a plain read-modify-write text replace; it is not a
      substitute for a real Word editing session.
    """
    original = paragraph.text
    count = original.count(find)
    if count == 0:
        return 0
    new_text = original.replace(find, replace)
    runs = paragraph.runs
    if runs:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ''
    return count


def _iter_all_paragraphs(doc):
    """Yield every paragraph in the document body and in every table cell."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


class IInstance(MicrosoftToolInstanceBase):
    IGlobal: IGlobal
    SERVICE = SERVICE

    # -----------------------------------------------------------------------
    # Helpers
    # -----------------------------------------------------------------------

    def _base(self) -> str:
        return graph_client.user_base(self.IGlobal.cfg)

    # =======================================================================
    # DIAGNOSTICS
    # =======================================================================

    @tool_function(
        description=(
            'Check the Word/Graph connection and verify that the granted OAuth scopes cover the '
            "node's configured access tier. Call this when a Word operation fails with a scope or "
            'permission error. Returns connection_ok: true when the required scopes are present.'
        ),
        input_schema={'type': 'object', 'properties': {}, 'required': []},
    )
    def word_check_connection(self, args: dict) -> dict:
        """Check the Word connection and whether granted OAuth scopes cover the access tier. Read-only."""
        base = self._base()

        def _probe(auth):
            request(auth, 'GET', f'{base}/drive')

        return self._check_connection_impl(probe=_probe)

    # =======================================================================
    # READ
    # =======================================================================

    @tool_function(
        input_schema={
            'type': 'object',
            'required': ['file'],
            'properties': {
                'file': {
                    'type': 'string',
                    'description': "Path in the acting user's OneDrive (e.g. 'Docs/report.docx') or a drive item id",
                },
            },
        },
        description=(
            'Read the text content of a .docx document: every body paragraph followed by every table '
            "cell's text, joined with newlines. Returns {text}."
        ),
    )
    def word_read_text(self, args: dict) -> dict:
        """Download a .docx and return its paragraph and table-cell text, newline-joined. Read-only."""
        import docx

        args = normalize_tool_input(args, tool_name='tool_word')
        file = require_str(args, 'file', tool_name='word_read_text')
        content, _etag = download_docx(self.IGlobal.auth, self._base(), file)
        doc = docx.Document(BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return {'text': '\n'.join(parts)}

    # =======================================================================
    # WRITE
    # =======================================================================

    @tool_function(
        input_schema={
            'type': 'object',
            'required': ['path', 'paragraphs'],
            'properties': {
                'path': {
                    'type': 'string',
                    'description': "Destination path in the acting user's OneDrive, e.g. 'Docs/report.docx'",
                },
                'paragraphs': {
                    'type': 'array',
                    'items': {'type': 'string'},
                    'description': 'Paragraph texts to add to the new document, in order',
                },
            },
        },
        description=(
            'Create a new .docx document at a OneDrive path from a list of paragraph texts, overwriting '
            'any existing file there. Returns the new file id, name, and webUrl. Requires the write tier.'
        ),
    )
    def word_create_document(self, args: dict) -> dict:
        """Create a new .docx document from paragraph texts at a OneDrive path. Requires the write tier."""
        args = normalize_tool_input(args, tool_name='tool_word')
        self.IGlobal.access.require_write('word_create_document')
        import docx

        path = require_str(args, 'path', tool_name='word_create_document')
        paragraphs = require_str_list(args, 'paragraphs', tool_name='word_create_document')
        doc = docx.Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        buf = BytesIO()
        doc.save(buf)
        # A brand-new file has no eTag to match — created directly rather
        # than through upload_docx (which always sends If-Match when given one).
        base = self._base()
        data = request(
            self.IGlobal.auth,
            'PUT',
            f'{base}/drive/root:/{urllib.parse.quote(path, safe=chr(47))}:/content',
            data=buf.getvalue(),
            content_type=WORD_CONTENT_TYPE,
        )
        return clean_item(data)

    @tool_function(
        input_schema={
            'type': 'object',
            'required': ['file', 'paragraphs'],
            'properties': {
                'file': {
                    'type': 'string',
                    'description': "Path in the acting user's OneDrive or a drive item id",
                },
                'paragraphs': {
                    'type': 'array',
                    'items': {'type': 'string'},
                    'description': 'Paragraph texts to append to the end of the document, in order',
                },
            },
        },
        description=(
            'Download a .docx, append paragraphs to the end, and re-upload with an If-Match precondition. '
            'Returns the updated file metadata. Requires the write tier. On a conflict (the file changed '
            'since it was read) the tool raises — re-read with word_read_text and retry.'
        ),
    )
    def word_append_text(self, args: dict) -> dict:
        """Append paragraphs to a .docx and re-upload with If-Match. Requires the write tier."""
        args = normalize_tool_input(args, tool_name='tool_word')
        self.IGlobal.access.require_write('word_append_text')
        import docx

        file = require_str(args, 'file', tool_name='word_append_text')
        paragraphs = require_str_list(args, 'paragraphs', tool_name='word_append_text')
        base = self._base()
        content, etag = download_docx(self.IGlobal.auth, base, file)
        doc = docx.Document(BytesIO(content))
        for text in paragraphs:
            doc.add_paragraph(text)
        buf = BytesIO()
        doc.save(buf)
        data = upload_docx(self.IGlobal.auth, base, file, buf.getvalue(), etag)
        return clean_item(data)

    @tool_function(
        input_schema={
            'type': 'object',
            'required': ['file', 'find', 'replace'],
            'properties': {
                'file': {
                    'type': 'string',
                    'description': "Path in the acting user's OneDrive or a drive item id",
                },
                'find': {'type': 'string', 'description': 'Text to find'},
                'replace': {'type': 'string', 'description': 'Replacement text'},
            },
        },
        description=(
            'Find and replace text throughout a .docx: every body paragraph and every table cell. Each '
            'matching paragraph is replaced as a whole (single pass over its original text, so a '
            'replacement that itself contains the search text is never re-matched). Paragraphs with no '
            'match are untouched; a paragraph made of multiple text runs loses its intra-paragraph '
            'formatting boundaries when it does match. Re-uploads with an If-Match precondition. Returns '
            '{replacements: <count>}. Requires the write tier.'
        ),
    )
    def word_replace_text(self, args: dict) -> dict:
        """Find/replace text across paragraphs and table cells; return the replacement count.

        Approach: see :func:`_replace_in_paragraph` — a single pass per
        paragraph over its *original* text (never over already-mutated
        text), so replacements aren't missed just because a paragraph's
        runs were split mid-word, and a ``replace`` value that itself
        contains ``find`` can never be re-matched and double-counted.
        Requires the write tier.
        """
        args = normalize_tool_input(args, tool_name='tool_word')
        self.IGlobal.access.require_write('word_replace_text')
        import docx

        file = require_str(args, 'file', tool_name='word_replace_text')
        find = require_str(args, 'find', tool_name='word_replace_text')
        replace = args.get('replace')
        if not isinstance(replace, str):
            raise ValueError('word_replace_text: "replace" is required and must be a string')
        base = self._base()
        content, etag = download_docx(self.IGlobal.auth, base, file)
        doc = docx.Document(BytesIO(content))
        total = sum(_replace_in_paragraph(p, find, replace) for p in _iter_all_paragraphs(doc))
        if total == 0:
            # Nothing matched: skip the upload so a no-op never bumps the item's eTag/mtime.
            return {'replacements': 0}
        buf = BytesIO()
        doc.save(buf)
        upload_docx(self.IGlobal.auth, base, file, buf.getvalue(), etag)
        return {'replacements': total}

    @tool_function(
        input_schema={
            'type': 'object',
            'required': ['file'],
            'properties': {
                'file': {
                    'type': 'string',
                    'description': "Path in the acting user's OneDrive or a drive item id",
                },
            },
        },
        description=(
            'Export a .docx to PDF via Graph server-side conversion and save the result beside the '
            "source file (same folder, source name with '.docx' stripped and '.pdf' appended). Returns "
            'the uploaded PDF file metadata — never the PDF bytes themselves, to avoid returning megabytes '
            'of content through the tool channel. Requires the write tier (it writes the new PDF file).'
        ),
    )
    def word_export_pdf(self, args: dict) -> dict:
        """Export a .docx to PDF and upload it beside the source. Requires the write tier."""
        args = normalize_tool_input(args, tool_name='tool_word')
        self.IGlobal.access.require_write('word_export_pdf')
        file = require_str(args, 'file', tool_name='word_export_pdf')
        base = self._base()
        auth = self.IGlobal.auth
        item = it(base, file)
        pdf_bytes = request(auth, 'GET', f'{item}/content', params={'format': 'pdf'}, binary=True)
        meta = request(auth, 'GET', item)
        name = meta.get('name') or file
        stem = name[:-5] if name.lower().endswith('.docx') else name
        pdf_name = f'{stem}.pdf'
        parent = meta.get('parentReference') or {}
        parent_path = parent.get('path') or '/drive/root:'
        dest = f'{base}{parent_path}/{_seg(pdf_name)}:/content'
        data = request(auth, 'PUT', dest, data=pdf_bytes, content_type='application/pdf')
        return clean_item(data)
