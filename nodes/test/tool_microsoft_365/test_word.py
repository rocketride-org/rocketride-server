# =============================================================================
# RocketRide Engine
# =============================================================================
# MIT License
# Copyright (c) 2026 Aparavi Software AG
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in
# all copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.
# =============================================================================

"""
Focused tier/round-trip tests for the Word service.

Real ``IInstance`` method calls with only the HTTP layer
(``graph_client._urlopen``) mocked, mirroring
``test_outlook_calendar.py``'s bootstrap. Tests that manipulate an actual
``.docx`` via ``python-docx`` are skipped (``pytest.importorskip``) when the
dependency is not installed in the current interpreter; the readonly-gate
test needs no docx content and always runs.
"""

from __future__ import annotations

import io
import json
import sys
import types
import urllib.error
from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path
from unittest import mock

import pytest

_TEST_DIR = Path(__file__).resolve().parents[2]  # nodes/test -> nodes
_REPO_ROOT = _TEST_DIR.parent
_NODES_SRC = _TEST_DIR / 'src'
if str(_NODES_SRC) not in sys.path:
    sys.path.insert(0, str(_NODES_SRC))

# Self-sufficient bootstrap (same technique as test_outlook_calendar.py /
# test_outlook_mail_guards.py): stub the engine runtime modules the word
# package imports (depends/rocketlib/ai.common.config), but load the *real*
# ai.common.utils.tool_args module directly by file path — it has no heavy
# deps (json/typing/rocketlib.warning only) — so normalize_tool_input/
# require_str/require_str_list behave exactly as in production instead of
# returning MagicMocks.
_added = []
for _name in ('depends', 'rocketlib', 'ai', 'ai.common', 'ai.common.config'):
    if _name not in sys.modules:
        _stub = types.ModuleType(_name)
        if _name == 'depends':
            _stub.depends = lambda *a, **k: None
        if _name == 'rocketlib':
            _stub.IInstanceBase = object
            _stub.IGlobalBase = object
            _stub.OPEN_MODE = types.SimpleNamespace(CONFIG='CONFIG')
            _stub.warning = lambda *a, **k: None
            _stub.tool_function = lambda **kw: lambda f: f
        if _name == 'ai.common.config':
            _stub.Config = object
        sys.modules[_name] = _stub
        _added.append(_name)

if 'ai.common.utils' not in sys.modules:
    _tool_args_path = _REPO_ROOT / 'packages' / 'ai' / 'src' / 'ai' / 'common' / 'utils' / 'tool_args.py'
    _spec = spec_from_file_location('ai.common.utils.tool_args', _tool_args_path)
    _tool_args = module_from_spec(_spec)
    sys.modules['ai.common.utils.tool_args'] = _tool_args
    _spec.loader.exec_module(_tool_args)
    _utils_mod = types.ModuleType('ai.common.utils')
    for _n in (
        'normalize_tool_input',
        'optional_str',
        'optional_str_list',
        'optional_bool',
        'require_str',
        'require_str_list',
        'int_arg',
    ):
        setattr(_utils_mod, _n, getattr(_tool_args, _n))
    sys.modules['ai.common.utils'] = _utils_mod
    _added.append('ai.common.utils')
    _added.append('ai.common.utils.tool_args')

_fresh_nodes = 'nodes' not in sys.modules
from nodes.tool_microsoft_365 import graph_client as gc  # noqa: E402
from nodes.tool_microsoft_365.word.IInstance import IInstance  # noqa: E402

from nodes.core.microsoft_access import MicrosoftAccessError, WORD, resolve_microsoft_access  # noqa: E402

for _name in _added:
    sys.modules.pop(_name, None)
if _fresh_nodes:
    for _name in [k for k in list(sys.modules) if k == 'nodes' or k.startswith('nodes.')]:
        sys.modules.pop(_name, None)

try:
    import docx

    HAS_DOCX = True
except ImportError:
    docx = None
    HAS_DOCX = False

requires_docx = pytest.mark.skipif(not HAS_DOCX, reason='python-docx is not installed')


def _json_resp(body: dict, status: int = 200, headers: dict | None = None):
    m = mock.MagicMock()
    m.read.return_value = json.dumps(body).encode()
    m.status = status
    m.headers = headers or {}
    m.__enter__ = lambda s: s
    m.__exit__ = lambda s, *a: False
    return m


def _binary_resp(raw: bytes, status: int = 200):
    m = mock.MagicMock()
    m.read.return_value = raw
    m.status = status
    m.headers = {}
    m.__enter__ = lambda s: s
    m.__exit__ = lambda s, *a: False
    return m


def _http_error(status: int, body: dict | None = None):
    payload = io.BytesIO(json.dumps({'error': body or {}}).encode())
    return urllib.error.HTTPError('u', status, 'err', {}, payload)


def _instance(*, tier: str) -> IInstance:
    inst = IInstance()
    cfg = {'access': tier}
    access = resolve_microsoft_access(cfg, WORD)
    auth = mock.MagicMock()
    auth.token.return_value = 'TOK'
    inst.IGlobal = types.SimpleNamespace(access=access, auth=auth, cfg={'authType': 'user'})
    return inst


def _build_docx_bytes(*, paragraph_text='hello world', cell_text='foo bar') -> bytes:
    doc = docx.Document()
    doc.add_paragraph(paragraph_text)
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_bytes_from(build) -> bytes:
    """Build a docx via a ``build(doc)`` callback (for paragraph/run shapes
    ``_build_docx_bytes`` can't express) and return its saved bytes.
    """
    doc = docx.Document()
    build(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestReadonlyBlocksWrite:
    # These block on IGlobal.access.require_write before any docx content is
    # touched, so they must run even when python-docx is not installed.
    def test_readonly_blocks_create_document(self):
        inst = _instance(tier='readonly')
        with pytest.raises(MicrosoftAccessError, match='read-only'):
            inst.word_create_document({'path': 'Docs/new.docx', 'paragraphs': ['hi']})

    def test_readonly_blocks_replace_text(self):
        inst = _instance(tier='readonly')
        with pytest.raises(MicrosoftAccessError, match='read-only'):
            inst.word_replace_text({'file': 'Docs/doc.docx', 'find': 'foo', 'replace': 'bar'})

    def test_readonly_blocks_append_text(self):
        inst = _instance(tier='readonly')
        with pytest.raises(MicrosoftAccessError, match='read-only'):
            inst.word_append_text({'file': 'Docs/doc.docx', 'paragraphs': ['x']})

    def test_readonly_blocks_export_pdf(self):
        inst = _instance(tier='readonly')
        with pytest.raises(MicrosoftAccessError, match='read-only'):
            inst.word_export_pdf({'file': 'Docs/doc.docx'})


@requires_docx
class TestCreateDocument:
    def test_create_document_builds_a_valid_docx(self):
        inst = _instance(tier='write')
        with mock.patch.object(
            gc, '_urlopen', return_value=_json_resp({'id': 'f1', 'name': 'new.docx', 'webUrl': 'https://x'})
        ) as u:
            out = inst.word_create_document({'path': 'Docs/new.docx', 'paragraphs': ['Hello', 'World']})
            assert out == {'id': 'f1', 'name': 'new.docx', 'webUrl': 'https://x'}
            req = u.call_args[0][0]
            assert req.get_method() == 'PUT'
            assert req.full_url.endswith('/drive/root:/Docs/new.docx:/content')
            assert req.get_header('Content-type') == (
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            # Round-trip: read the paragraphs back out of the PUT body bytes.
            uploaded = docx.Document(io.BytesIO(req.data))
            assert [p.text for p in uploaded.paragraphs] == ['Hello', 'World']
            # No etag on a brand-new file: no If-Match header sent.
            assert req.get_header('If-match') is None


@requires_docx
class TestReplaceText:
    def test_replace_text_counts_paragraph_and_table_cell(self):
        inst = _instance(tier='write')
        original = _build_docx_bytes(paragraph_text='foo is here', cell_text='another foo cell')
        meta_resp = _json_resp({'eTag': '"abc"'})
        content_resp = _binary_resp(original)
        put_resp = _json_resp({'id': 'f1', 'name': 'doc.docx'})
        with mock.patch.object(gc, '_urlopen', side_effect=[meta_resp, content_resp, put_resp]) as u:
            out = inst.word_replace_text({'file': 'Docs/doc.docx', 'find': 'foo', 'replace': 'bar'})
            assert out == {'replacements': 2}
            put_req = u.call_args_list[2][0][0]
            assert put_req.get_method() == 'PUT'
            assert put_req.get_header('If-match') == '"abc"'
            uploaded = docx.Document(io.BytesIO(put_req.data))
            assert uploaded.paragraphs[0].text == 'bar is here'
            assert uploaded.tables[0].rows[0].cells[0].text == 'another bar cell'

    def test_replace_text_with_zero_matches_does_not_upload(self):
        # Regression: a no-op replace must not PUT the document (would bump eTag/mtime).
        inst = _instance(tier='write')
        original = _build_docx_bytes(paragraph_text='nothing here', cell_text='nor here')
        meta_resp = _json_resp({'eTag': '"abc"'})
        content_resp = _binary_resp(original)
        with mock.patch.object(gc, '_urlopen', side_effect=[meta_resp, content_resp]) as u:
            out = inst.word_replace_text({'file': 'Docs/doc.docx', 'find': 'foo', 'replace': 'bar'})
            assert out == {'replacements': 0}
            assert u.call_count == 2
            assert all(c[0][0].get_method() == 'GET' for c in u.call_args_list)

    def test_replace_where_replacement_contains_find_does_not_double_count_or_corrupt(self):
        # Regression: a single-pass replace scanning ONLY the original text.
        # A run-then-paragraph two-pass approach would re-scan the just-written
        # replacement, matching 'foo' again inside 'foobar' and corrupting the
        # result to 'foobarbar is here' with a wrong count of 2.
        inst = _instance(tier='write')
        original = _build_docx_bytes(paragraph_text='foo is here', cell_text='no match here')
        meta_resp = _json_resp({'eTag': '"abc"'})
        content_resp = _binary_resp(original)
        put_resp = _json_resp({'id': 'f1', 'name': 'doc.docx'})
        with mock.patch.object(gc, '_urlopen', side_effect=[meta_resp, content_resp, put_resp]) as u:
            out = inst.word_replace_text({'file': 'Docs/doc.docx', 'find': 'foo', 'replace': 'foobar'})
            assert out == {'replacements': 1}
            put_req = u.call_args_list[2][0][0]
            uploaded = docx.Document(io.BytesIO(put_req.data))
            assert uploaded.paragraphs[0].text == 'foobar is here'

    def test_replace_across_a_run_boundary(self):
        # 'TODO' split into two separate runs ('TO' + 'DO') must still be
        # found and replaced as a whole, with a correct count of 1.
        def _build(doc):
            paragraph = doc.add_paragraph()
            paragraph.add_run('TO')
            paragraph.add_run('DO')
            paragraph.add_run(': fix this')

        original = _docx_bytes_from(_build)
        inst = _instance(tier='write')
        meta_resp = _json_resp({'eTag': '"abc"'})
        content_resp = _binary_resp(original)
        put_resp = _json_resp({'id': 'f1', 'name': 'doc.docx'})
        with mock.patch.object(gc, '_urlopen', side_effect=[meta_resp, content_resp, put_resp]) as u:
            out = inst.word_replace_text({'file': 'Docs/doc.docx', 'find': 'TODO', 'replace': 'DONE'})
            assert out == {'replacements': 1}
            put_req = u.call_args_list[2][0][0]
            uploaded = docx.Document(io.BytesIO(put_req.data))
            assert uploaded.paragraphs[0].text == 'DONE: fix this'

    def test_replace_two_matches_in_one_paragraph_counts_both(self):
        inst = _instance(tier='write')
        original = _build_docx_bytes(paragraph_text='foo and foo again', cell_text='no match here')
        meta_resp = _json_resp({'eTag': '"abc"'})
        content_resp = _binary_resp(original)
        put_resp = _json_resp({'id': 'f1', 'name': 'doc.docx'})
        with mock.patch.object(gc, '_urlopen', side_effect=[meta_resp, content_resp, put_resp]) as u:
            out = inst.word_replace_text({'file': 'Docs/doc.docx', 'find': 'foo', 'replace': 'bar'})
            assert out == {'replacements': 2}
            put_req = u.call_args_list[2][0][0]
            uploaded = docx.Document(io.BytesIO(put_req.data))
            assert uploaded.paragraphs[0].text == 'bar and bar again'


@requires_docx
class TestReadText:
    def test_read_text_joins_paragraphs_and_table_text(self):
        inst = _instance(tier='readonly')
        original = _build_docx_bytes(paragraph_text='line one', cell_text='cell text')
        meta_resp = _json_resp({'eTag': '"abc"'})
        content_resp = _binary_resp(original)
        with mock.patch.object(gc, '_urlopen', side_effect=[meta_resp, content_resp]):
            out = inst.word_read_text({'file': 'Docs/doc.docx'})
            assert out == {'text': 'line one\ncell text'}


@requires_docx
class TestAppendText:
    def test_append_text_adds_paragraphs_and_sends_if_match(self):
        inst = _instance(tier='write')
        original = _build_docx_bytes(paragraph_text='first')
        meta_resp = _json_resp({'eTag': '"xyz"'})
        content_resp = _binary_resp(original)
        put_resp = _json_resp({'id': 'f1', 'name': 'doc.docx'})
        with mock.patch.object(gc, '_urlopen', side_effect=[meta_resp, content_resp, put_resp]) as u:
            inst.word_append_text({'file': 'Docs/doc.docx', 'paragraphs': ['second']})
            put_req = u.call_args_list[2][0][0]
            assert put_req.get_header('If-match') == '"xyz"'
            uploaded = docx.Document(io.BytesIO(put_req.data))
            assert [p.text for p in uploaded.paragraphs][-1] == 'second'


@requires_docx
class TestConflict:
    def test_append_text_conflict_surfaces_as_graph_error(self):
        inst = _instance(tier='write')
        original = _build_docx_bytes()
        meta_resp = _json_resp({'eTag': '"stale"'})
        content_resp = _binary_resp(original)
        with mock.patch.object(gc, '_urlopen', side_effect=[meta_resp, content_resp, _http_error(412)]):
            with pytest.raises(gc.GraphError, match='conflict'):
                inst.word_append_text({'file': 'Docs/doc.docx', 'paragraphs': ['x']})
